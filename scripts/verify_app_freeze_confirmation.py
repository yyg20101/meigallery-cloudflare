#!/usr/bin/env python3
"""校验 App 1.0 需求冻结准备清单与客户确认 DOCX。"""

from __future__ import annotations

import hashlib
import importlib.util
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
GENERATOR = ROOT / "scripts/generate_app_freeze_confirmation.py"


def load_generator():
    spec = importlib.util.spec_from_file_location(
        "generate_app_freeze_confirmation",
        GENERATOR,
    )
    if spec is None or spec.loader is None:
        raise RuntimeError("无法加载需求冻结确认单生成器")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


gen = load_generator()


def fail(message: str) -> None:
    raise SystemExit(f"需求冻结确认资料校验失败：{message}")


def document_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def verify_source_fingerprints(markdown: str) -> str:
    current = gen.source_fingerprints()
    expected_combined = gen.combined_fingerprint(current)
    match = re.search(r"^基线指纹：`([0-9a-f]{64})`$", markdown, re.MULTILINE)
    if not match:
        fail("Markdown 缺少完整组合基线指纹")
    if match.group(1) != expected_combined:
        fail(
            "Markdown 组合基线指纹已过期；"
            f"记录={match.group(1)}，当前={expected_combined}"
        )

    for path, digest in current:
        pattern = (
            rf"^\|\s*{re.escape(path)}\s*\|\s*{digest}\s*\|\s*"
            rf"{digest[:12]}\s*\|$"
        )
        if not re.search(pattern, markdown, re.MULTILINE):
            fail(f"Markdown 缺少或错配冻结文件指纹：{path}")
    return expected_combined


def verify_markdown(markdown: str) -> str:
    fingerprint = verify_source_fingerprints(markdown)
    required = (
        "状态：冻结准备中，尚未授权开发",
        "92 个 Page ID",
        "349 个 Figma 最终状态",
        "2,284 个有效交互动作",
        "169 个确定性图片映射",
        "全量像素级视觉设计已完成",
        "功能/交互冻结",
        "像素级视觉冻结",
        "不创建 KMP 工程",
        "不新增 App API",
        "不执行数据库 migration",
    )
    for text in required:
        if text not in markdown:
            fail(f"Markdown 缺少关键冻结边界：{text}")

    for index in range(1, 9):
        decision_id = f"C-{index:02d}"
        if decision_id not in markdown:
            fail(f"Markdown 缺少客户决策：{decision_id}")
    for index in range(1, 8):
        gate_id = f"G-{index:02d}"
        if gate_id not in markdown:
            fail(f"Markdown 缺少专业门禁：{gate_id}")

    if "状态：App 1.0 需求已冻结" in markdown:
        fail("客户尚未签署，Markdown 不得提前标记为已冻结")
    return fingerprint


def verify_docx(fingerprint: str) -> tuple[int, int]:
    if not gen.DOCX_OUTPUT.exists():
        fail(f"DOCX 不存在：{gen.DOCX_OUTPUT}")

    with zipfile.ZipFile(gen.DOCX_OUTPUT) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            fail(f"DOCX 压缩包损坏：{bad_member}")
        document_xml = archive.read("word/document.xml")

    doc = Document(gen.DOCX_OUTPUT)
    text = document_text(doc)
    required = (
        "冻结准备中",
        "尚未授权开发",
        "92 个 Page ID",
        "349 个 Figma 最终状态",
        "2,284 个有效交互动作",
        "169/169 已校验",
        "已完成，待签署",
        fingerprint,
    )
    for value in required:
        if value not in text:
            fail(f"DOCX 缺少关键冻结内容：{value}")

    for item in gen.CUSTOMER_DECISIONS:
        if item.decision_id not in text or item.title not in text:
            fail(f"DOCX 缺少客户决策：{item.decision_id}")
    for gate in gen.PROFESSIONAL_GATES:
        if gate.gate_id not in text or gate.title not in text:
            fail(f"DOCX 缺少专业门禁：{gate.gate_id}")

    for path, digest in gen.source_fingerprints():
        if Path(path).name not in text or digest[:12] not in text:
            fail(f"DOCX 缺少冻结文件短指纹：{path}")

    if "App 1.0 需求已冻结" in text and "才能改为“App 1.0 需求已冻结”" not in text:
        fail("客户尚未签署，DOCX 不得提前标记为已冻结")

    image_count = len(doc.inline_shapes)
    if image_count != len(gen.VISUAL_EVIDENCE):
        fail(
            f"DOCX 内嵌图片={image_count}，"
            f"应为 {len(gen.VISUAL_EVIDENCE)}"
        )
    alt_count = len(re.findall(rb"<wp:docPr\b[^>]*\bdescr=", document_xml))
    if alt_count != image_count:
        fail(f"DOCX 图片替代文本={alt_count}，应为 {image_count}")
    for alt, _ in gen.VISUAL_EVIDENCE:
        if alt.encode("utf-8") not in document_xml:
            fail(f"DOCX 缺少图片替代文本：{alt}")

    for table_index, table in enumerate(doc.tables, start=1):
        if not table.rows:
            fail(f"DOCX 表格 {table_index} 没有行")
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            fail(f"DOCX 表格 {table_index} 首行未标记为表头")
        tbl_pr = table._tbl.tblPr
        width = tbl_pr.find(qn("w:tblW"))
        indent = tbl_pr.find(qn("w:tblInd"))
        if width is None or width.get(qn("w:w")) != "9360":
            fail(f"DOCX 表格 {table_index} 宽度不是 9360 DXA")
        if indent is None or indent.get(qn("w:w")) != "120":
            fail(f"DOCX 表格 {table_index} 缩进不是 120 DXA")

    section = doc.sections[0]
    if section.page_width.inches != 8.5 or section.page_height.inches != 11:
        fail("DOCX 页面尺寸不是 US Letter")
    margins = (
        section.top_margin.inches,
        section.right_margin.inches,
        section.bottom_margin.inches,
        section.left_margin.inches,
    )
    if any(abs(value - 1.0) > 0.01 for value in margins):
        fail(f"DOCX 页边距不是 1 英寸：{margins}")

    return image_count, len(doc.tables)


def verify_manifest_counts() -> None:
    manifest = json.loads(gen.PAGE_MANIFEST.read_text(encoding="utf-8"))
    counts = manifest.get("counts", {})
    expected = {
        "pages": 92,
        "totalCaptures": 146,
        "detailedFigmaPages": 5,
        "detailedFigmaStateCaptures": 23,
        "documentPrototypeMappings": 169,
        "figmaDesignedPages": 92,
        "figmaDesignedStates": 349,
        "figmaFlowPreviews": 92,
        "figmaTotalActions": 2284,
    }
    for key, value in expected.items():
        if counts.get(key) != value:
            fail(f"原型清单 {key}={counts.get(key)!r}，应为 {value}")


def main() -> None:
    if not gen.MD_OUTPUT.exists():
        fail(f"Markdown 不存在：{gen.MD_OUTPUT}")
    markdown = gen.MD_OUTPUT.read_text(encoding="utf-8")
    fingerprint = verify_markdown(markdown)
    verify_manifest_counts()
    image_count, table_count = verify_docx(fingerprint)
    digest = hashlib.sha256(gen.DOCX_OUTPUT.read_bytes()).hexdigest()
    print(
        "需求冻结确认资料校验通过："
        f"客户决策={len(gen.CUSTOMER_DECISIONS)}，"
        f"专业门禁={len(gen.PROFESSIONAL_GATES)}，"
        f"DOCX 表格={table_count}，图片={image_count}，"
        f"基线指纹={fingerprint[:12]}，"
        f"DOCX 指纹={digest[:12]}。"
    )


if __name__ == "__main__":
    main()
