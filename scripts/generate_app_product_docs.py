#!/usr/bin/env python3
"""从 App 产品 Markdown 基线生成客户可评审的 DOCX 交付物。

版式约束：
- 基础预设：standard_business_brief。
- 首页模板：customer_pack。
- 命名覆盖：标题与标题层级使用 MeiGallery 品牌粉色；为保证中文在
  Word、LibreOffice 与 macOS 预览中均可读，正文统一使用 Arial Unicode MS，
  缺少该字体的办公环境由系统中文字体替代。
  正文尺寸、间距、列表缩进、表格几何仍严格沿用预设。
"""

from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from PIL import Image as PILImage


ROOT = Path(__file__).resolve().parents[1]


def resolve_document_skill_scripts() -> Path:
    base = (
        Path.home()
        / ".codex/plugins/cache/openai-primary-runtime/documents"
    )
    candidates = sorted(
        base.glob("*/skills/documents/scripts"),
        reverse=True,
    )
    for candidate in candidates:
        if (candidate / "table_geometry.py").exists():
            return candidate
    raise FileNotFoundError("未找到 documents skill 的 table_geometry.py")


DOC_SKILL_SCRIPTS = resolve_document_skill_scripts()
sys.path.insert(0, str(DOC_SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


DELIVERABLES = ROOT / "docs/app/deliverables"
PAGE_ASSET_ROOT = ROOT / "docs/app/assets/page-prototypes"
PAGE_MANIFEST = PAGE_ASSET_ROOT / "manifest.json"
BODY_FONT = "Arial Unicode MS"
EAST_ASIA_FONT = "Arial Unicode MS"
MONO_FONT = "JetBrains Mono"
BRAND = RGBColor(216, 61, 115)
BRAND_DARK = RGBColor(143, 36, 75)
INK = RGBColor(31, 31, 34)
MUTED = RGBColor(104, 105, 112)
LIGHT_FILL = "F2F4F7"
PINK_FILL = "FFF1F6"
PINK_LINE = "E7A3BC"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


@dataclass(frozen=True)
class DocumentSpec:
    source: Path
    output: Path
    title: str
    subtitle: str
    running_label: str
    purpose: str
    scope: str
    review_path: tuple[str, ...]
    supplemental_images: tuple[tuple[str, Path], ...] = ()


SPECS = (
    DocumentSpec(
        source=ROOT / "docs/app/MEIGALLERY_APP_1_0_CLIENT_PRD.md",
        output=DELIVERABLES / "MeiGallery_App_1.0_产品需求确认书.docx",
        title="MeiGallery App 1.0\n产品需求确认书",
        subtitle="客户需求确认、产品设计、开发估算与测试验收的单一业务基线",
        running_label="MeiGallery App 1.0｜产品需求确认书",
        purpose="用于客户确认 App 1.0 范围、角色、业务规则、交互、运营边界和验收标准。",
        scope="独立 Android/iOS App、Nuxt 管理后台、现有 MeiGallery 数据复用与未来迁移。",
        review_path=(
            "先阅读第 1、3、7、8、9、16 章，确认产品边界与核心业务闭环。",
            "再操作高保真原型，核对发现、会员申请、平台话题、运营回复、会员发放和金币调整。",
            "最后填写第 17 章客户确认事项与第 19 章确认结论。",
        ),
    ),
    DocumentSpec(
        source=ROOT / "docs/app/APP_PAGE_LEVEL_PRODUCT_DESIGN.md",
        output=DELIVERABLES / "MeiGallery_App_1.0_逐页交互设计确认册.docx",
        title="MeiGallery App 1.0\n逐页交互设计确认册",
        subtitle="移动端 50 页 + 管理后台 49 页，共 99 个页面级设计对象",
        running_label="MeiGallery App 1.0｜逐页交互设计确认册",
        purpose="用于按 Page ID 逐页确认页面目标、入口、操作、异常状态、出口和验收条件。",
        scope="Android/iOS 移动端与 Nuxt 管理后台；99 页是完整需求覆盖，不等于首批同时开发。",
        review_path=(
            "先确认全局交互规则与不可变产品边界。",
            "按 P0 关键旅程评审移动端，再评审后台闭环。",
            "逐页意见必须引用 Page ID，并注明具体状态和预期调整。",
        ),
    ),
)


INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
CHECK_RE = re.compile(r"^\s*[-*]\s+\[([ xX])\]\s+(.+)$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.+)$")
NUMBER_RE = re.compile(r"^(\s*)\d+\.\s+(.+)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*:?-{3,}:?\s*$")


def set_run_font(
    run,
    *,
    latin: str = BODY_FONT,
    east_asia: str = EAST_ASIA_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color: str = "D8DCE3", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{side}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_bottom_rule(paragraph, color: str = PINK_LINE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "7")
    bottom.set(qn("w:color"), color)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "686970")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(color)
    r_pr.append(size)
    field_run.append(r_pr)
    field_run.append(fld_begin)
    field_run.append(instr)
    field_run.append(fld_separate)
    field_run.append(fallback)
    field_run.append(fld_end)
    paragraph._p.append(field_run)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def set_link_style(run) -> None:
    set_run_font(run, color=BRAND_DARK)
    run.font.underline = True


def add_hyperlink(paragraph, text: str, target: str) -> None:
    relation_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    run = paragraph.add_run(text)
    set_link_style(run)
    new_run.append(run._r.get_or_add_rPr())
    for child in list(run._r):
        if child.tag != qn("w:rPr"):
            new_run.append(child)
    paragraph._p.remove(run._r)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_content(
    paragraph,
    text: str,
    *,
    base_size: float = 11,
    base_color: RGBColor = INK,
    base_bold: bool = False,
) -> None:
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, color=base_color, bold=base_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                latin=MONO_FONT,
                east_asia=EAST_ASIA_FONT,
                size=max(8.5, base_size - 0.8),
                color=BRAND_DARK,
            )
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), PINK_FILL)
            run._r.get_or_add_rPr().append(shading)
        else:
            link_match = re.match(r"^\[([^\]]+)\]\(([^)]+)\)$", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color, bold=base_bold)


def configure_page(section) -> None:
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BRAND, 16, 8),
        ("Heading 2", 13, BRAND, 12, 6),
        ("Heading 3", 12, BRAND_DARK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True


def next_abstract_num_id(numbering) -> int:
    ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId"))
    ]
    return max(ids, default=-1) + 1


def next_num_id(numbering) -> int:
    ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    ]
    return max(ids, default=0) + 1


def create_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    marker = "%1." if kind == "decimal" else ("☐" if kind == "check" else "•")
    lvl_text.set(qn("w:val"), marker)
    level.append(lvl_text)
    justify = OxmlElement("w:lvlJc")
    justify.set(qn("w:val"), "left")
    level.append(justify)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num_id = next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def configure_header_footer(section, label: str) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    set_run_font(run, size=9, color=MUTED, bold=True)
    add_bottom_rule(paragraph, color="E2C1CD", size="5")

    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    add_page_field(paragraph)


def add_cover(doc: Document, spec: DocumentSpec) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("APP 1.0 · 客户确认资料")
    set_run_font(run, size=10.5, color=BRAND, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    for index, line in enumerate(spec.title.split("\n")):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=30, color=BRAND_DARK, bold=True)
    add_bottom_rule(p, color="E5A3BA", size="12")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(spec.subtitle)
    set_run_font(run, size=13.5, color=MUTED)

    table = doc.add_table(rows=2, cols=4)
    values = (
        ("产品", "MeiGallery App", "版本", "1.0"),
        ("状态", "待客户确认", "日期", "2026-07-30"),
    )
    for row_index, row_values in enumerate(values):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, PINK_FILL if col_index % 2 == 0 else WHITE)
            set_cell_border(cell, color="E7C4D1")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(
                run,
                size=9.5,
                color=BRAND_DARK if col_index % 2 == 0 else INK,
                bold=col_index % 2 == 0,
            )
    mark_repeat_header(table.rows[0])
    apply_table_geometry(
        table,
        column_widths_from_weights((0.85, 2.25, 0.85, 2.55), TABLE_WIDTH_DXA),
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("文档用途")
    set_run_font(run, size=11, color=BRAND, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline_content(p, spec.purpose)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("覆盖范围")
    set_run_font(run, size=11, color=BRAND, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline_content(p, spec.scope)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    add_inline_content(
        p,
        "说明：本文中的具体数值均为“建议基线，待客户确认”；"
        "客户未确认前，不视为已承诺的商业指标或最终研发参数。",
        base_size=10,
        base_color=BRAND_DARK,
        base_bold=True,
    )
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PINK_FILL)
    p_pr.append(shading)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "180")
    p_pr.append(ind)

    doc.add_page_break()
    heading = doc.add_paragraph("评审阅读路径", style="Heading 1")
    set_keep_with_next(heading)
    num_id = create_numbering(doc, "decimal")
    for item in spec.review_path:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        add_inline_content(p, item)

    heading = doc.add_paragraph("交付物关系", style="Heading 2")
    set_keep_with_next(heading)
    for item in (
        "产品需求确认书：确认产品范围、业务闭环、业务规则与验收。",
        "逐页交互设计确认册：按 Page ID 确认每一页的目标、状态与操作。",
        "可交互原型：现场演示关键旅程和 99 页独立页面设计。",
    ):
        p = doc.add_paragraph()
        add_inline_content(p, item)

    add_revision_note(doc, style="Heading 2")

    if spec.supplemental_images:
        heading = doc.add_paragraph("原型总览", style="Heading 2")
        set_keep_with_next(heading)
        if len(spec.supplemental_images) == 2:
            add_image_pair(doc, spec.supplemental_images)
        else:
            for alt, image_path in spec.supplemental_images:
                add_image(doc, alt, image_path, width=Inches(6.35))

    doc.add_page_break()


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", stripped)]


def is_table_block(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    if not lines[index].lstrip().startswith("|"):
        return False
    if not lines[index + 1].lstrip().startswith("|"):
        return False
    separator_cells = parse_table_row(lines[index + 1])
    return bool(separator_cells) and all(
        TABLE_SEPARATOR_RE.match(cell) for cell in separator_cells
    )


def table_weights(headers: list[str], rows: list[list[str]]) -> list[float]:
    header_text = " ".join(headers)
    if "Page ID" in header_text and len(headers) == 5:
        return [1.45, 1.2, 2.6, 1.45, 2.0]
    if "心遇" in header_text and len(headers) >= 5:
        return [2.0] + [1.0] * (len(headers) - 1)
    if headers[:2] == ["编号", "确认事项"] and len(headers) == 4:
        return [0.65, 1.65, 3.35, 1.25]
    if len(headers) == 2:
        return [1.6, 4.9]
    if len(headers) == 3:
        return [1.55, 2.5, 2.45]
    if len(headers) == 4:
        return [1.15, 1.8, 2.65, 1.4]
    if len(headers) == 5:
        return [1.15, 1.45, 2.45, 1.45, 2.0]
    if len(headers) == 6:
        return [1.35, 1.0, 1.0, 1.0, 1.0, 1.0]

    def visible_length(value: str) -> int:
        value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
        value = re.sub(r"[`*_]", "", value)
        return max(4, min(32, len(value)))

    weights: list[float] = []
    for col in range(len(headers)):
        lengths = [visible_length(headers[col])]
        lengths.extend(
            visible_length(row[col]) for row in rows if col < len(row)
        )
        weights.append(float(min(22, max(7, max(lengths)))))
    return weights


def add_markdown_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    mark_repeat_header(header_row)
    prevent_row_split(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_inline_content(
            p,
            header,
            base_size=9,
            base_color=BRAND_DARK,
            base_bold=True,
        )

    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        normalized = row_values + [""] * (len(headers) - len(row_values))
        for index, value in enumerate(normalized[: len(headers)]):
            cell = row.cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, WHITE)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            if len(headers) >= 5 and index in (0, 1, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_content(p, value, base_size=8.6, base_color=INK)

    widths = column_widths_from_weights(
        table_weights(headers, rows),
        TABLE_WIDTH_DXA,
    )
    apply_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def set_image_alt(inline_shape, alt: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt)


def fitted_picture_width(
    path: Path,
    *,
    max_width,
    max_height,
):
    """按原图比例同时约束宽高，避免竖版 Figma 截图跨页裁切。"""
    with PILImage.open(path) as image:
        width_px, height_px = image.size
    if width_px <= 0 or height_px <= 0:
        return max_width
    width_inches = min(
        max_width.inches,
        max_height.inches * width_px / height_px,
    )
    return Inches(width_inches)


def add_image(
    doc: Document,
    alt: str,
    path: Path,
    *,
    width=Inches(6.35),
    max_height=Inches(5.8),
) -> None:
    if not path.exists():
        p = doc.add_paragraph()
        add_inline_content(
            p,
            f"原型图片暂不可用：{alt}",
            base_color=BRAND_DARK,
            base_bold=True,
        )
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(
        str(path),
        width=fitted_picture_width(
            path,
            max_width=width,
            max_height=max_height,
        ),
    )
    set_image_alt(shape, alt)
    caption = doc.add_paragraph(f"图：{alt}", style="Caption")
    caption.paragraph_format.keep_with_next = False


def add_image_pair(
    doc: Document,
    images: tuple[tuple[str, Path], tuple[str, Path]],
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    for index, (alt, path) in enumerate(images):
        if not path.exists():
            continue
        if index:
            spacer = p.add_run("   ")
            set_run_font(spacer, size=6, color=MUTED)
        run = p.add_run()
        shape = run.add_picture(str(path), width=Inches(3.02))
        set_image_alt(shape, alt)
    caption = doc.add_paragraph(
        f"图：{images[0][0]}与{images[1][0]}",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = False


def add_code_block(doc: Document, content: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shading)
    run = p.add_run(content.rstrip())
    set_run_font(
        run,
        latin=MONO_FONT,
        east_asia=EAST_ASIA_FONT,
        size=8.6,
        color=RGBColor(53, 55, 60),
    )


def markdown_body_lines(source: Path) -> list[str]:
    lines = source.read_text(encoding="utf-8").splitlines()
    for index, line in enumerate(lines):
        if line.startswith("## "):
            return lines[index:]
    return lines


def render_markdown(doc: Document, source: Path) -> None:
    lines = markdown_body_lines(source)
    bullet_id = create_numbering(doc, "bullet")
    check_id = create_numbering(doc, "check")
    decimal_id: int | None = None
    previous_was_number = False
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            previous_was_number = False
            index += 1
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            previous_was_number = False
            markdown_level = len(heading_match.group(1))
            if markdown_level == 1:
                index += 1
                continue
            style_level = min(3, markdown_level - 1)
            p = doc.add_paragraph(heading_match.group(2), style=f"Heading {style_level}")
            set_keep_with_next(p)
            index += 1
            continue

        image_match = IMAGE_RE.match(line)
        if image_match:
            previous_was_number = False
            alt, relative_path = image_match.groups()
            image_path = (source.parent / relative_path).resolve()
            add_image(doc, alt or image_path.stem, image_path)
            index += 1
            continue

        if line.startswith("```"):
            previous_was_number = False
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, "\n".join(code_lines))
            index += 1
            continue

        if is_table_block(lines, index):
            previous_was_number = False
            headers = parse_table_row(lines[index])
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            add_markdown_table(doc, headers, rows)
            continue

        check_match = CHECK_RE.match(line)
        if check_match:
            previous_was_number = False
            p = doc.add_paragraph()
            apply_numbering(p, check_id)
            add_inline_content(p, check_match.group(2))
            index += 1
            continue

        number_match = NUMBER_RE.match(line)
        if number_match:
            if not previous_was_number or decimal_id is None:
                decimal_id = create_numbering(doc, "decimal")
            p = doc.add_paragraph()
            apply_numbering(p, decimal_id)
            add_inline_content(p, number_match.group(2))
            previous_was_number = True
            index += 1
            continue

        bullet_match = BULLET_RE.match(line)
        if bullet_match:
            previous_was_number = False
            p = doc.add_paragraph()
            apply_numbering(p, bullet_id)
            add_inline_content(p, bullet_match.group(2))
            index += 1
            continue

        if line == "---":
            previous_was_number = False
            p = doc.add_paragraph()
            add_bottom_rule(p, color="D9DCE2", size="4")
            index += 1
            continue

        if line.startswith(">"):
            previous_was_number = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            add_inline_content(p, line.lstrip("> ").strip(), base_color=MUTED)
            index += 1
            continue

        p = doc.add_paragraph()
        previous_was_number = False
        add_inline_content(p, line)
        index += 1


def add_revision_note(doc: Document, *, style: str = "Heading 1") -> None:
    heading = doc.add_paragraph("文档维护说明", style=style)
    set_keep_with_next(heading)
    p = doc.add_paragraph()
    add_inline_content(
        p,
        "本文与 App 版本保持一致，当前统一为 1.0。需求讨论期间直接更新同一基线，"
        "不为每次讨论创建无意义的小版本或兼容章节；已确认后的范围变更需先完成影响评估。",
    )


def load_page_manifest() -> dict[str, Any]:
    if not PAGE_MANIFEST.exists():
        raise FileNotFoundError(f"逐页原型清单不存在：{PAGE_MANIFEST}")
    manifest = json.loads(PAGE_MANIFEST.read_text(encoding="utf-8"))
    counts = manifest.get("counts", {})
    expected = {
        "pages": 99,
        "mobilePages": 50,
        "adminPages": 49,
        "p0Pages": 57,
        "p1Pages": 32,
        "p2Pages": 10,
        "defaultCaptures": 99,
        "keyStateCaptures": 57,
        "totalCaptures": 156,
        "detailedFigmaPages": 5,
        "detailedFigmaStateCaptures": 23,
        "documentPrototypeMappings": 179,
        "figmaDesignedPages": 99,
        "figmaDesignedStates": 408,
        "figmaMobileStates": 208,
        "figmaAdminStates": 200,
        "figmaFlowPreviews": 99,
        "figmaHistoricalPageActionBaseline": 2957,
        "figmaHistoricalFlowActionBaseline": 614,
        "figmaHistoricalActionBaseline": 3571,
    }
    for key, value in expected.items():
        if counts.get(key) != value:
            raise ValueError(
                f"逐页原型清单计数异常：{key}={counts.get(key)!r}，应为 {value}"
            )
    pages = manifest.get("pages", [])
    state_totals = {"mobile": 0, "admin": 0}
    for page in pages:
        states = page.get("states")
        if not isinstance(states, list) or not states:
            raise ValueError(f"{page.get('pageId', '未知页面')} 缺少正式状态")
        if any(
            not isinstance(state, str) or not state or state.strip() != state
            for state in states
        ):
            raise ValueError(
                f"{page.get('pageId', '未知页面')} 存在空白或未规范化的正式状态名称"
            )
        if len(states) != len(set(states)):
            raise ValueError(f"{page['pageId']} 存在重复正式状态")
        platform = page.get("platform")
        if platform not in state_totals:
            raise ValueError(f"{page['pageId']} 平台类型无效：{platform!r}")
        state_totals[platform] += len(states)
    derived = {
        "figmaDesignedPages": len(pages),
        "figmaDesignedStates": sum(state_totals.values()),
        "figmaMobileStates": state_totals["mobile"],
        "figmaAdminStates": state_totals["admin"],
    }
    for key, value in derived.items():
        if counts.get(key) != value:
            raise ValueError(
                f"逐页原型清单 {key} 与页面状态明细不一致："
                f"清单 {counts.get(key)!r}，实际 {value}"
            )
    if manifest.get("status") != "verified":
        raise ValueError("逐页原型清单尚未完成校验，拒绝生成客户文档")
    if int(manifest.get("schemaVersion", 0)) < 4:
        raise ValueError("逐页原型清单缺少 Figma 全量最终交付与需求追踪 schema")
    for page in pages:
        requirements = page.get("requirements", {})
        if not requirements.get("traceKey"):
            raise ValueError(f"{page.get('pageId', '未知页面')} 缺少需求追踪键")
        if not requirements.get("product") or not requirements.get("release"):
            raise ValueError(f"{page['pageId']} 缺少产品或发布范围需求编号")
        if not requirements.get("features"):
            raise ValueError(f"{page['pageId']} 缺少 Feature PRD 映射")
    return manifest


def add_compact_field(
    doc: Document,
    label: str,
    value: str,
    *,
    size: float = 8.35,
    keep_together: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = keep_together
    run = paragraph.add_run(f"{label}｜")
    set_run_font(run, size=size, color=BRAND_DARK, bold=True)
    run = paragraph.add_run(value)
    set_run_font(run, size=size, color=INK)


def add_confirmation_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    add_inline_content(
        paragraph,
        f"客户确认｜□ 接受  □ 调整  □ 不适用    意见：{text}",
        base_size=8.5,
        base_color=BRAND_DARK,
        base_bold=True,
    )
    add_bottom_rule(paragraph, color="D9A8BA", size="4")


def capture_maps(
    manifest: dict[str, Any],
) -> tuple[dict[str, dict[str, Any]], dict[str, dict[str, Any]]]:
    defaults: dict[str, dict[str, Any]] = {}
    key_states: dict[str, dict[str, Any]] = {}
    for capture in manifest["captures"]:
        target = defaults if capture["variant"] == "default" else key_states
        target[capture["pageId"]] = capture
    return defaults, key_states


def figma_capture_map(
    manifest: dict[str, Any],
) -> dict[tuple[str, str], dict[str, Any]]:
    captures = manifest.get("figmaStateCaptures", [])
    expected = manifest["counts"]["detailedFigmaStateCaptures"]
    if len(captures) != expected:
        raise ValueError(
            f"Figma 逐状态导出图数量异常：{len(captures)}，应为 {expected}"
        )
    result: dict[tuple[str, str], dict[str, Any]] = {}
    for capture in captures:
        key = (capture["pageId"], capture["state"])
        if key in result:
            raise ValueError(
                f"Figma 逐状态导出映射重复：{capture['pageId']} / {capture['state']}"
            )
        result[key] = capture
    return result


def preferred_capture(
    capture: dict[str, Any],
    final_states: dict[tuple[str, str], dict[str, Any]],
) -> dict[str, Any]:
    """用已审计的 Figma 最终图替换同 Page ID、同状态的基础占位图。

    基础映射的 alt 和 variant 继续保留，以便 156 个基础图片映射不失效；
    图片、Frame 和交互链接来自最终状态，避免客户看到过时原型。
    """
    final = final_states.get((capture["pageId"], capture["state"]))
    if final is None:
        return capture
    merged = dict(capture)
    for key in (
        "image",
        "expectedWidth",
        "expectedHeight",
        "width",
        "height",
        "sha256",
        "bytes",
        "sourceUrl",
        "prototypeUrl",
        "frameId",
        "screen",
    ):
        if key in final:
            merged[key] = final[key]
    merged["usesFigmaFinalState"] = True
    return merged


def add_page_confirmation_unit(
    doc: Document,
    page: dict[str, Any],
    capture: dict[str, Any],
) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.page_break_before = True
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    kicker.paragraph_format.keep_with_next = True
    add_inline_content(
        kicker,
        f"{'移动端' if page['platform'] == 'mobile' else '管理后台'} · "
        f"{page['module']} · {page['priority']} · 逐页确认单元",
        base_size=8.8,
        base_color=BRAND,
        base_bold=True,
    )

    heading = doc.add_paragraph(
        f"{page['pageId']}  {page['pageName']}",
        style="Heading 1",
    )
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(3)

    purpose = doc.add_paragraph()
    purpose.paragraph_format.space_after = Pt(2)
    add_inline_content(
        purpose,
        page["purpose"],
        base_size=9.2,
        base_color=INK,
        base_bold=True,
    )

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(2)
    add_inline_content(
        metadata,
        f"设计路由：`{page['route']}`　默认状态：{capture['state']}　"
        f"角色：{page['roles']}　Figma：{page['figmaDesignPage']} / "
        f"{page['figmaDesignedStateCount']} 个状态",
        base_size=8.4,
        base_color=MUTED,
    )

    add_image(
        doc,
        capture["alt"],
        PAGE_ASSET_ROOT / capture["image"],
        width=Inches(6.15),
        max_height=Inches(4.3),
    )

    add_compact_field(
        doc,
        "入口与前置",
        f"{page['entry']}。{page['preconditions']}",
    )
    add_compact_field(
        doc,
        "页面结构与交互",
        f"{page['structure']} {page['interaction']}",
    )
    add_compact_field(
        doc,
        "业务规则与权限",
        f"{page['rule']} {page['dataPermission']}",
    )
    add_compact_field(
        doc,
        "需求追踪",
        page["requirements"]["traceKey"],
        size=7.75,
        keep_together=True,
    )
    add_compact_field(
        doc,
        "状态与下一步",
        f"{'、'.join(page['states'])}；成功后可进入 "
        f"{page['nextPageId'] or '当前流程安全出口'}。",
        keep_together=True,
    )
    add_compact_field(
        doc,
        "验收要点",
        "；".join(page["acceptance"]),
        size=8.15,
    )
    add_confirmation_line(doc, "________________________________")


def add_page_catalog_appendix(
    doc: Document,
    manifest: dict[str, Any],
) -> None:
    defaults, _ = capture_maps(manifest)
    final_states = figma_capture_map(manifest)
    heading = doc.add_paragraph("附录 A：99 页详细需求与默认原型", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    set_keep_with_next(heading)
    paragraph = doc.add_paragraph()
    add_inline_content(
        paragraph,
        "本附录以 Page ID 为唯一映射键。每个页面均包含用途、入口、角色、"
        "结构、交互、规则、权限、需求追踪、状态、验收标准、对应默认原型和客户确认栏；"
        "移动端 50 页、管理后台 49 页，共 99 个独立确认单元。",
        base_size=10,
    )
    paragraph = doc.add_paragraph()
    add_inline_content(
        paragraph,
        "评审规则：原型中的 Page ID、页面名称和状态必须与本页标题及说明一致；"
        "任何调整意见均应同时写明 Page ID 和目标状态。",
        base_size=9.5,
        base_color=BRAND_DARK,
        base_bold=True,
    )

    pages = sorted(manifest["pages"], key=lambda item: item["order"])
    for page in pages:
        capture = defaults.get(page["pageId"])
        if capture is None:
            raise ValueError(f"缺少默认原型：{page['pageId']}")
        add_page_confirmation_unit(
            doc,
            page,
            preferred_capture(capture, final_states),
        )


def add_key_state_image(
    doc: Document,
    page: dict[str, Any],
    capture: dict[str, Any],
) -> None:
    image_path = PAGE_ASSET_ROOT / capture["image"]
    if not image_path.exists():
        raise FileNotFoundError(f"关键状态原型不存在：{image_path}")

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    add_inline_content(
        title,
        f"{page['pageId']} {page['pageName']}｜关键状态：{capture['state']}",
        base_size=8.7,
        base_color=BRAND_DARK,
        base_bold=True,
    )

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(0)
    image_paragraph.paragraph_format.keep_with_next = True
    shape = image_paragraph.add_run().add_picture(
        str(image_path),
        width=fitted_picture_width(
            image_path,
            max_width=Inches(6.05),
            max_height=Inches(3.75),
        ),
    )
    set_image_alt(shape, capture["alt"])

    caption = doc.add_paragraph(
        f"图：{capture['alt']}；触发后必须提供可理解原因和安全下一步。",
        style="Caption",
    )
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = False


def add_key_state_appendix(
    doc: Document,
    manifest: dict[str, Any],
) -> None:
    _, key_states = capture_maps(manifest)
    final_states = figma_capture_map(manifest)
    pages_by_id = {page["pageId"]: page for page in manifest["pages"]}
    captures = [
        capture
        for capture in manifest["captures"]
        if capture["variant"] == "key-state"
    ]
    if len(captures) != manifest["counts"]["keyStateCaptures"]:
        raise ValueError(
            f"关键状态截图数量异常：{len(captures)}，"
            f"应为 {manifest['counts']['keyStateCaptures']}"
        )
    captures.sort(key=lambda item: pages_by_id[item["pageId"]]["order"])

    heading = doc.add_paragraph("附录 B：57 个 P0 关键状态原型", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    set_keep_with_next(heading)
    paragraph = doc.add_paragraph()
    add_inline_content(
        paragraph,
        "P0 页面除默认状态外，另提供一个对开发与验收最关键的异常、受限、"
        "冲突或完成状态。以下 56 张原型与附录 A 使用同一 Page ID 映射。",
        base_size=10,
    )

    for index in range(0, len(captures), 2):
        group = doc.add_paragraph(
            f"P0 关键状态 {index + 1:02d}–{min(index + 2, len(captures)):02d}",
            style="Heading 2",
        )
        group.paragraph_format.page_break_before = True
        group.paragraph_format.space_before = Pt(0)
        group.paragraph_format.space_after = Pt(4)
        for capture in captures[index : index + 2]:
            page = pages_by_id[capture["pageId"]]
            if key_states.get(page["pageId"]) != capture:
                raise ValueError(f"关键状态映射冲突：{page['pageId']}")
            add_key_state_image(
                doc,
                page,
                preferred_capture(capture, final_states),
            )


def add_detailed_figma_state_page(
    doc: Document,
    page: dict[str, Any],
    capture: dict[str, Any],
) -> None:
    image_path = PAGE_ASSET_ROOT / capture["image"]
    if not image_path.exists():
        raise FileNotFoundError(f"Figma 最终状态原型不存在：{image_path}")

    kicker = doc.add_paragraph()
    kicker.paragraph_format.page_break_before = True
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    kicker.paragraph_format.keep_with_next = True
    add_inline_content(
        kicker,
        f"Figma 最终交互 · 状态 {capture['stateIndex']}/{capture['stateCount']} · "
        f"Frame `{capture['frameId']}`",
        base_size=8.5,
        base_color=BRAND,
        base_bold=True,
    )

    heading = doc.add_paragraph(
        f"{page['pageId']}  {page['pageName']}｜{capture['state']}",
        style="Heading 1",
    )
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(2)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(2)
    add_inline_content(
        metadata,
        f"页面优先级：{page['priority']}　适用角色：{page['roles']}",
        base_size=8.2,
        base_color=MUTED,
    )

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(1)
    shape = image_paragraph.add_run().add_picture(
        str(image_path),
        width=Inches(2.15),
    )
    set_image_alt(shape, capture["alt"])

    caption = doc.add_paragraph(
        f"图：{capture['screen']}；本地导出与 Figma Frame 一一对应。",
        style="Caption",
    )
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(2)

    add_compact_field(
        doc,
        "触发条件",
        capture["trigger"],
        size=7.75,
    )
    add_compact_field(
        doc,
        "关键交互",
        capture["interaction"],
        size=7.75,
    )
    add_compact_field(
        doc,
        "预期结果",
        capture["expected"],
        size=7.75,
    )
    add_compact_field(
        doc,
        "权威边界",
        capture["authority"],
        size=7.75,
    )

    link = doc.add_paragraph()
    link.paragraph_format.space_before = Pt(1)
    link.paragraph_format.space_after = Pt(2)
    add_inline_content(
        link,
        f"[打开 Figma 交互原型]({capture['prototypeUrl']})",
        base_size=8.1,
        base_color=BRAND_DARK,
        base_bold=True,
    )
    add_confirmation_line(doc, "________________________________")


def add_detailed_figma_state_appendix(
    doc: Document,
    manifest: dict[str, Any],
) -> None:
    captures = manifest.get("figmaStateCaptures", [])
    pages_by_id = {page["pageId"]: page for page in manifest["pages"]}
    page_ids = list(dict.fromkeys(capture["pageId"] for capture in captures))
    if len(page_ids) != manifest["counts"]["detailedFigmaPages"]:
        raise ValueError(
            f"Figma 最终细化页面数量异常：{len(page_ids)}，"
            f"应为 {manifest['counts']['detailedFigmaPages']}"
        )

    heading = doc.add_paragraph(
        "附录 C：通知与金币 23 张 Figma 逐状态导出图",
        style="Heading 1",
    )
    heading.paragraph_format.page_break_before = True
    set_keep_with_next(heading)
    paragraph = doc.add_paragraph()
    add_inline_content(
        paragraph,
        "Figma 最终设计已覆盖 99 页、408 个需求状态；本附录另外保留通知列表、"
        "通知详情、金币钱包、金币明细和金币分录详情 5 个页面的 23 张逐状态本地导出图。"
        "每个导出状态均绑定 Page ID、状态名、"
        "Figma Frame ID、触发条件、关键交互、预期结果与服务端权威边界。",
        base_size=10,
    )
    paragraph = doc.add_paragraph()
    add_inline_content(
        paragraph,
        "最终设计覆盖移动端 50 页/208 状态、后台 49 页/200 状态；"
        "3,571 个页面内与流程交互源、缺失目标 0、移动端不足 44dp 的关键热区 0 和"
        "文字溢出 0 均为 APP-SET-08 增量六态前历史审计基线，新增六态已完成定向 QA，"
        "当前全量动作与 QA 统计待开发结束后统一重算。附录 A/B 中这 5 个页面的同名状态已自动使用逐状态导出图。",
        base_size=9.5,
        base_color=BRAND_DARK,
        base_bold=True,
    )

    for capture in captures:
        page = pages_by_id.get(capture["pageId"])
        if page is None:
            raise ValueError(f"Figma 最终状态缺少页面：{capture['pageId']}")
        add_detailed_figma_state_page(doc, page, capture)


def add_final_delivery_confirmation(
    doc: Document,
    manifest: dict[str, Any],
) -> None:
    counts = manifest["counts"]
    heading = doc.add_paragraph("逐页需求与原型总确认", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    set_keep_with_next(heading)
    for text in (
        f"本文件已纳入 {counts['pages']} 个页面默认原型，其中移动端 "
        f"{counts['mobilePages']} 页、管理后台 {counts['adminPages']} 页。",
        f"P0 页面共 {counts['p0Pages']} 个，并分别纳入一个关键状态原型；"
        f"默认原型与关键状态合计 {counts['totalCaptures']} 张。",
        f"Figma 最终设计覆盖 {counts['figmaDesignedPages']} 页、"
        f"{counts['figmaDesignedStates']} 个需求状态和 "
        f"APP-SET-08 增量前历史基线为 {counts['figmaHistoricalActionBaseline']:,} 个有效交互动作，"
        "当前动作总数待全部开发完成后统一重算；"
        f"移动端/后台状态分别为 {counts['figmaMobileStates']}/"
        f"{counts['figmaAdminStates']}。",
        f"通知与金币 {counts['detailedFigmaPages']} 个页面另保留 "
        f"{counts['detailedFigmaStateCaptures']} 张逐状态本地导出图；"
        f"客户文档共建立 {counts['documentPrototypeMappings']} 个确定性图片映射。",
        f"Figma 最终版本 ID 为 {manifest['figmaFinal']['finalVersionId']}；"
        "开发和评审均应在 `40｜Delivery Index` 中按 Page ID 定位。",
        "每个 Page ID 均包含产品总需求编号、App 1.0 发布范围编号和 "
        "Feature PRD 需求组组成的追踪键。",
        "每张图片均通过 Page ID、页面名称、状态和文件清单建立确定性映射；"
        "客户意见应引用 Page ID，避免口头描述造成页面错配。",
    ):
        paragraph = doc.add_paragraph()
        add_inline_content(paragraph, text, base_size=10.5)

    table = doc.add_table(rows=4, cols=2)
    mark_repeat_header(table.rows[0])
    values = (
        ("□ 全部确认", "范围、逐页需求、原型和关键状态可作为后续实现与验收基线。"),
        ("□ 有条件确认", "除下方明确列出的 Page ID 外，其余内容确认。"),
        ("□ 暂不确认", "需重新评审的 Page ID 与原因见下方意见。"),
        ("客户意见", "\n\n"),
    )
    for row, values_row in zip(table.rows, values):
        prevent_row_split(row)
        for index, value in enumerate(values_row):
            cell = row.cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, PINK_FILL if index == 0 else WHITE)
            set_cell_border(cell, color="E7C4D1")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_content(
                paragraph,
                value,
                base_size=9.2,
                base_color=BRAND_DARK if index == 0 else INK,
                base_bold=index == 0,
            )
    apply_table_geometry(
        table,
        column_widths_from_weights((1.55, 4.95), TABLE_WIDTH_DXA),
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(20)
    add_inline_content(
        paragraph,
        "客户代表签字：____________________    日期：____年__月__日",
        base_size=10.5,
        base_bold=True,
    )


def build_document(spec: DocumentSpec) -> None:
    manifest = load_page_manifest()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_page(section)
    configure_header_footer(section, spec.running_label)
    add_cover(doc, spec)
    render_markdown(doc, spec.source)
    add_page_catalog_appendix(doc, manifest)
    add_key_state_appendix(doc, manifest)
    add_detailed_figma_state_appendix(doc, manifest)
    add_final_delivery_confirmation(doc, manifest)

    doc.core_properties.title = spec.title.replace("\n", " ")
    doc.core_properties.subject = spec.subtitle
    doc.core_properties.author = "MeiGallery 产品团队"
    doc.core_properties.keywords = "MeiGallery, App 1.0, PRD, 交互设计, 客户确认"
    doc.core_properties.comments = (
        "版式：standard_business_brief；首页：customer_pack；"
        "命名覆盖：MeiGallery 品牌粉色标题层级、Arial Unicode MS 中文字体；"
        "逐页原型：99 张默认状态 + 57 张 P0 关键状态；"
        "Figma 最终设计：99 页、408 个状态；3,571 个有效交互源为增量前历史基线，待开发结束后统一重算；"
        "客户文档图片映射：179 个。"
    )
    doc.core_properties.last_modified_by = "MeiGallery 产品团队"

    spec.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(spec.output)
    print(f"已生成：{spec.output}")


def main() -> None:
    for spec in SPECS:
        build_document(spec)


if __name__ == "__main__":
    main()
