#!/usr/bin/env python3
"""生成 App 1.0 历史需求冻结准备快照（Markdown + 客户确认 DOCX）。

文档定位：
- Markdown 保留冻结准备阶段的客户决策、门禁与变更控制历史快照。
- DOCX 是客户签署产品范围、交互基线和调整意见的短版确认单。
- 本脚本不把“待确认”自动改为“已冻结”，客户和专业责任人签署后才能执行冻结。

版式：
- 基础预设：compact_reference_guide。
- 首页模板：customer_pack。
- 命名覆盖：沿用 MeiGallery 品牌粉色与 Hiragino Sans GB 中文字体。
"""

from __future__ import annotations

import hashlib
import importlib.util
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
APP_ROOT = ROOT / "docs/app"
DELIVERABLES = APP_ROOT / "deliverables"
MD_OUTPUT = APP_ROOT / "APP_1_0_REQUIREMENTS_FREEZE_CHECKLIST.md"
DOCX_OUTPUT = DELIVERABLES / "MeiGallery_App_1.0_需求冻结确认单.docx"
PAGE_MANIFEST = APP_ROOT / "assets/page-prototypes/manifest.json"


def load_product_doc_helpers():
    module_path = ROOT / "scripts/generate_app_product_docs.py"
    spec = importlib.util.spec_from_file_location(
        "generate_app_product_docs",
        module_path,
    )
    if spec is None or spec.loader is None:
        raise RuntimeError("无法加载产品 DOCX 生成辅助模块")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


pd = load_product_doc_helpers()


@dataclass(frozen=True)
class CustomerDecision:
    decision_id: str
    title: str
    recommendation: str
    impact: str
    deadline: str
    source_ids: str


@dataclass(frozen=True)
class ProfessionalGate:
    gate_id: str
    title: str
    conclusion: str
    owner: str
    deadline: str
    source_ids: str


CUSTOMER_DECISIONS = (
    CustomerDecision(
        "C-01",
        "产品定位与当前沟通主体",
        "定位为“经授权真人内容发现与平台话题服务”；App 1.0 的话题由平台运营接收。",
        "决定商店描述、详情页披露、会员价值和投诉边界。",
        "客户需求确认书签署前",
        "OQ-034",
    ),
    CustomerDecision(
        "C-02",
        "用户沟通入口",
        "当前统一使用“发起话题（平台接收）”；仅未来完成本人认领和复核后才显示“给本人发私信”。",
        "决定详情、会员、会话和通知中的身份文案。",
        "客户需求确认书签署前",
        "OQ-034",
    ),
    CustomerDecision(
        "C-03",
        "五级心享会员权益",
        "每日新话题额度采用 1/2/4/6/10；保存筛选采用 1/3/6/12/20；收藏夹采用 3/5/10/20/30。",
        "决定 entitlement、运营容量和会员页展示。",
        "entitlement API 冻结前",
        "OQ-014",
    ),
    CustomerDecision(
        "C-04",
        "会员获取闭环",
        "用户提交站内申请，管理员审核并手动发放；通常在 1 个服务日内首次处理；App 1.0 不提供购买。",
        "决定会员申请、后台队列、通知和客服口径。",
        "会员流程开发前",
        "OQ-010、OQ-014",
    ),
    CustomerDecision(
        "C-05",
        "平台运营服务承诺",
        "服务时段每日 10:00–22:00；不保证固定回复；队列满载时暂停创建新话题，历史仍可读。",
        "决定排班、容量降级、页面状态和用户预期。",
        "平台话题 Alpha 前",
        "OQ-010、OQ-035",
    ),
    CustomerDecision(
        "C-06",
        "首发供给与发布方式",
        "至少 80 个合格真人资料、3 个地区组，以限量邀请 Beta 首发；公开扩量另行决策。",
        "决定内容准备成本、测试样本和推荐可信度。",
        "Beta 邀请前",
        "OQ-035",
    ),
    CustomerDecision(
        "C-07",
        "研发优先级",
        "采用 P0/P1/P2 分层：P0 阻塞 Alpha/Beta，P1 阻塞 App 1.0 最终验收，P2 独立排期。",
        "决定估算、里程碑和“99 页设计覆盖”与“首批开发范围”的区别。",
        "研发估算前",
        "OQ-036",
    ),
    CustomerDecision(
        "C-08",
        "正式产品名与品牌",
        "“心动遇见你”继续作为工作名；完成商标、域名和商店重名检索后确定正式名称。",
        "不阻塞功能范围确认，但阻塞应用 ID、商店物料和最终视觉品牌冻结。",
        "客户端脚手架与商店物料前",
        "OQ-001",
    ),
)


PROFESSIONAL_GATES = (
    ProfessionalGate(
        "G-01",
        "首发地区、运营主体、年龄与登录",
        "书面确定首发地区/语言、法律运营主体、成年人门槛、登录方式和目标渠道。",
        "Owner、产品、法务、平台",
        "账号与协议 schema 冻结前",
        "OQ-002、OQ-003、OQ-030",
    ),
    ProfessionalGate(
        "G-02",
        "真人来源、用途授权与认证范围",
        "完成现有 MeiGallery 授权盘点，冻结 App 用途授权、认证声明、发布复核和双人分离规则。",
        "内容 Owner、法务、安全、产品",
        "首批导入与认证实现前",
        "OQ-006、OQ-007、OQ-008",
    ),
    ProfessionalGate(
        "G-03",
        "数据保留、数据位置与隐私责任",
        "分别确认账号、消息、授权证据和审计保留期，以及 Cloudflare 数据位置、跨境结论和隐私 Owner。",
        "隐私、法务、安全、Cloudflare 负责人",
        "数据库 schema 与生产敏感数据接入前",
        "OQ-020、OQ-024、OQ-025",
    ),
    ProfessionalGate(
        "G-04",
        "消息审核、举报 SLA 与撤回证据",
        "冻结文本审核组合、P0/P1/P2 安全 SLA、7×24 升级联系人、撤回窗口和关闭后重开规则。",
        "安全负责人、运营、隐私、产品",
        "平台话题 API 与 Alpha 前",
        "OQ-021、OQ-022、OQ-033",
    ),
    ProfessionalGate(
        "G-05",
        "会员迁移与金币内控参数",
        "确认旧 vip/svip 迁移规则、调币阈值、负余额限制、批量与双人复核规则。",
        "Owner、产品、财务、安全",
        "权益迁移与财务后台实现前",
        "OQ-016、OQ-018",
    ),
    ProfessionalGate(
        "G-06",
        "KMP 构建矩阵与系统最低版本",
        "通过最小 Android/iOS 工程验证后，冻结 Kotlin、Compose、Gradle、AGP、KSP、JDK、Xcode、iOS 最低版本及 Android 是否高于 API 26。",
        "客户端架构负责人、产品",
        "客户端脚手架前",
        "OQ-026、OQ-027",
    ),
    ProfessionalGate(
        "G-07",
        "实时容量、Cloudflare 预算与运营排班",
        "验证消息实时恢复、容量与灾备，确认 Cloudflare 预算、服务时段排班、备岗和队列降级阈值。",
        "Owner、后端、SRE、运营",
        "Alpha 压测与 Beta 邀请前",
        "OQ-028、OQ-031、OQ-035",
    ),
)


SOURCE_FILES = (
    APP_ROOT / "MEIGALLERY_APP_1_0_CLIENT_PRD.md",
    APP_ROOT / "MEIGALLERY_APP_1_0_DEVELOPMENT_REQUIREMENTS.md",
    APP_ROOT / "APP_REQUIREMENTS_TRACEABILITY.md",
    PAGE_MANIFEST,
    DELIVERABLES / "MeiGallery_App_1.0_产品需求确认书.docx",
    DELIVERABLES / "MeiGallery_App_1.0_逐页交互设计确认册.docx",
)


VISUAL_EVIDENCE = (
    (
        "移动端发现与原始视觉参考对照",
        APP_ROOT
        / "assets/figma-qa/phase3/comparison-discovery-reference-vs-official-20260730.png",
    ),
    (
        "移动端五级会员最终设计",
        APP_ROOT
        / "assets/figma-qa/phase3/mobile-membership-catalog-20260730.png",
    ),
    (
        "管理后台真人工作台最终设计",
        APP_ROOT
        / "assets/figma-qa/phase3/admin-person-workbench-normal-20260730.jpeg",
    ),
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def source_fingerprints() -> list[tuple[str, str]]:
    result = []
    for path in SOURCE_FILES:
        if not path.exists():
            raise FileNotFoundError(f"冻结基线文件不存在：{path}")
        result.append((path.relative_to(ROOT).as_posix(), sha256(path)))
    return result


def combined_fingerprint(fingerprints: Iterable[tuple[str, str]]) -> str:
    payload = "\n".join(f"{path}:{digest}" for path, digest in fingerprints)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def load_manifest() -> dict:
    data = json.loads(PAGE_MANIFEST.read_text(encoding="utf-8"))
    expected = {
        "pages": 99,
        "mobilePages": 50,
        "adminPages": 49,
        "p0Pages": 57,
        "p1Pages": 32,
        "p2Pages": 10,
        "totalCaptures": 156,
        "detailedFigmaPages": 5,
        "detailedFigmaStateCaptures": 23,
        "documentPrototypeMappings": 179,
        "figmaDesignedPages": 99,
        "figmaDesignedStates": 408,
        "figmaMobileStates": 208,
        "figmaAdminStates": 200,
        "figmaFlowPreviews": 99,
        "figmaHistoricalActionBaseline": 3571,
    }
    for key, value in expected.items():
        if data.get("counts", {}).get(key) != value:
            raise ValueError(
                f"冻结清单拒绝生成：manifest {key}="
                f"{data.get('counts', {}).get(key)!r}，应为 {value}"
            )
    if data.get("status") != "verified":
        raise ValueError("冻结清单拒绝生成：逐页原型清单尚未验证")
    return data


def markdown_table(headers: tuple[str, ...], rows: Iterable[tuple[str, ...]]) -> str:
    output = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join("---" for _ in headers) + "|",
    ]
    for row in rows:
        output.append(
            "| "
            + " | ".join(value.replace("\n", "<br>") for value in row)
            + " |"
        )
    return "\n".join(output)


def generate_markdown(
    fingerprints: list[tuple[str, str]],
    fingerprint: str,
) -> None:
    decision_rows = tuple(
        (
            item.decision_id,
            item.title,
            item.recommendation,
            item.impact,
            item.deadline,
            item.source_ids,
            "□ 接受　□ 调整　□ 暂缓",
        )
        for item in CUSTOMER_DECISIONS
    )
    gate_rows = tuple(
        (
            item.gate_id,
            item.title,
            item.conclusion,
            item.owner,
            item.deadline,
            item.source_ids,
            "待责任人书面关闭",
        )
        for item in PROFESSIONAL_GATES
    )
    fingerprint_rows = tuple(
        (path, digest, digest[:12]) for path, digest in fingerprints
    )

    text = f"""# MeiGallery App 1.0 需求冻结准备与确认清单

App 版本：1.0

日期：2026-07-30

状态：历史快照（2026-07-30 冻结准备阶段，尚未授权开发）

当前事实提示：截至 2026-08-24，Figma 为 99 页/408 个正式状态，移动端 50 页/208 状态、后台 49 页/200 状态，当前动作总数为 3,585；本快照正文中的 402/202 与 3,571 个动作均为 `APP-SET-08` 六态增量前历史基线。

基线指纹：`{fingerprint}`

## 1. 结论

当前资料已具备客户确认产品范围、功能规则和页面交互的条件，但不应提前标记为“需求已冻结”：

- 产品与业务：8 项客户决策尚待逐项接受、调整或明确暂缓。
- 页面与交互：99 个 Page ID、402 个 Figma 最终状态、99 个流程预览和 3,571 个有效交互动作已完成；客户文档映射基线已建立 179 个确定性图片映射，可进入逐页确认。
- 视觉成熟度：移动端 50 页/202 状态、管理后台 49 页/200 状态均已完成像素级 Figma 最终稿和 QA；是否冻结仍等待客户/设计负责人签署。
- 专业门禁：7 组法律、隐私、安全、运营、财务和技术结论尚需责任人关闭；这些门禁不替客户决定功能，但会阻止对应 schema、工程或生产发布。
- 开发授权：在客户签署和阻塞门禁关闭前，不创建 KMP 工程、不新增 App API、不执行数据库 migration。

## 2. 冻结对象与成熟度

{markdown_table(
    ("对象", "当前证据", "当前结论", "冻结条件"),
    (
        ("产品范围", "客户 PRD 第 1–19 章、C-01 至 C-08", "可确认，未冻结", "8 项客户决策形成书面结论"),
        ("功能需求", "41 个产品需求编号、99 个追踪键", "可确认，未冻结", "客户范围结论同步到开发规格和追踪矩阵"),
        ("页面交互", "50 个移动端页、49 个后台页", "可逐页确认，未冻结", "逐页确认册意见关闭或形成明确例外"),
        ("原型映射", "156 张基础原型 + 23 张逐状态导出", "179/179 已建立", "Page ID、状态、图片和需求追踪继续保持一致"),
        ("像素级视觉", "99 页 / 402 状态 Figma 最终稿", "已完成，待签署", "客户与设计负责人按 Page ID 关闭意见"),
        ("技术与上线", "架构、KMP、API、Cloudflare 和运营文档", "规划完整，门禁未关闭", "G-01 至 G-07 按最晚关闭点完成"),
    ),
)}

## 3. 客户需逐项确认的 8 项决策

填写规则：每项必须选择“接受”“调整”或“暂缓”。选择调整时必须写明替代结论；选择暂缓时必须接受对应的阻塞范围和最晚关闭点。

{markdown_table(
    ("ID", "确认事项", "建议基线", "影响", "最晚关闭点", "来源", "客户选择"),
    decision_rows,
)}

客户调整意见：

1. `C-__`：____________________________________________________________
2. `C-__`：____________________________________________________________
3. `C-__`：____________________________________________________________

## 4. 上线前专业门禁

客户签署产品范围不等同于替法务、安全、财务、运营或技术负责人作专业结论。下列门禁必须由责任方分别提供书面证据：

{markdown_table(
    ("Gate", "门禁", "必须形成的结论", "责任方", "最晚关闭点", "来源", "状态"),
    gate_rows,
)}

## 5. 交互与视觉冻结说明

### 5.1 已具备确认条件

- 99 个 Page ID 均有页面目标、角色、入口、结构、主次操作、状态、权限、出口和验收条件。
- 57 个 P0 页面除默认状态外均提供一个关键异常、受限、冲突或处理中状态。
- Figma 最终文件完整覆盖 99 个 Page ID、402 个需求状态、99 个流程预览和 3,571 个有效交互动作。
- 客户 DOCX 下一次生成将保留 156 张基础逐页图与通知/金币 23 张逐状态导出图，共 179 个确定性图片映射。

### 5.2 不得混淆的边界

- “功能/交互冻结”表示信息结构、业务规则、权限、状态、文案含义和跨页出口不再被下游人员自行改变。
- “像素级视觉冻结”表示颜色、字体、间距、图片裁切、组件形态和动效参数已经获得设计确认。
- 全量像素级视觉设计已完成，但只有客户/设计负责人签署后才能把“设计已完成”改为“视觉已冻结”。
- 管理后台 UI 要求低于用户 App，但权限、状态、审批、并发冲突和审计交互仍必须按逐页规格验收。

### 5.3 最终视觉确认方法

1. 先通过 `40｜Delivery Index` 按 Page ID 核对 99 页和 402 个状态。
2. 再通过 `30｜Prototype Flows` 核对 99 个流程的主操作、异常状态和安全出口。
3. 客户意见必须引用 Page ID、状态和目标调整；只描述截图位置的意见不进入冻结清单。
4. 修改后重新执行文字、Icon、热区、溢出、交互目标和产品边界 QA。

## 6. 冻结生效条件

仅当以下条件全部满足时，状态才能从“冻结准备中”改为“App 1.0 需求已冻结”：

1. C-01 至 C-07 已选择并签署；C-08 已确认正式名称，或书面接受“工作名继续使用并阻塞品牌相关工程”的处理。
2. 产品需求确认书的总确认结论已签署。
3. 逐页交互设计确认册的修改意见已关闭，或形成带 Page ID、状态、责任人和关闭日期的例外清单。
4. 客户调整已同步到客户 PRD、开发需求规格、产品决策基线、开放问题、Feature PRD、追踪矩阵、页面目录、原型和 DOCX。
5. 99 页、402 个 Figma 最终状态、3,571 个有效动作和 179 个客户文档图片映射的自动校验继续通过。
6. 阻塞下一阶段的专业门禁已按 G-01 至 G-07 的最晚关闭点完成。
7. 产品 Owner、客户代表、设计负责人和技术负责人共同确认冻结结论。

## 7. 冻结后的变更控制

冻结后仍保持 App 与文档版本为 1.0，不因讨论次数递增文档小版本。任何实质变化必须先提交变更说明，至少包含：

- 变更原因、客户价值和不变更的后果。
- 受影响的 C/G/OQ、产品需求编号、Feature PRD、Page ID、状态和原型。
- 对 API、数据、权限、安全、运营、测试、工期和存量数据的影响。
- 是否需要重新客户确认、重新视觉确认或正常发布新 App 版本。
- 决策人、决定日期、生效日期和回滚条件。

纯文字修正、错别字或不改变含义的排版调整可以直接更新；变化历史由 Git 保存，不在正文中堆叠兼容章节。

## 8. 客户确认结论

- [ ] 同意 C-01 至 C-08 的建议基线，进入需求冻结。
- [ ] 原则同意，但需按本清单的客户调整意见同步后再冻结。
- [ ] 只同意产品范围与功能交互；Figma 最终视觉需按所列 Page ID 修改后再次确认。
- [ ] 暂不确认，需要重新讨论产品定位、范围或服务承诺。

客户/项目方：____________________________

确认人：__________________________________

签字/盖章：________________________________

确认日期：________年____月____日

## 9. 冻结包清单与指纹

以下文件共同构成本次待确认基线。任一文件内容变化后必须重新生成本清单，旧签署件不得自动覆盖新内容。

{markdown_table(
    ("基线文件", "SHA-256", "短指纹"),
    fingerprint_rows,
)}

## 10. 签署后的执行顺序

1. 将客户选择和调整同步到全部上游与下游文档。
2. 重新生成两份完整客户 DOCX、逐页原型清单和本确认单。
3. 运行需求、Page ID、原型、DOCX、无障碍和全页渲染校验。
4. 由 Owner 将文档状态改为“App 1.0 需求已冻结”并记录冻结指纹。
5. 关闭客户视觉意见和阻塞工程门禁，再进入 API/DTO 数据契约冻结与 KMP 脚手架。
"""
    MD_OUTPUT.write_text(text, encoding="utf-8")


def configure_compact_styles(doc: Document) -> None:
    """应用 compact_reference_guide，并使用 MeiGallery 品牌命名覆盖。"""
    pd.configure_styles(doc)
    styles = doc.styles

    normal = styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, pd.BRAND, 18, 10),
        ("Heading 2", 13, pd.BRAND, 14, 7),
        ("Heading 3", 12, pd.BRAND_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0


def configure_customer_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("MeiGallery App 1.0｜需求冻结确认单")
    pd.set_run_font(run, size=9, color=pd.MUTED, bold=True)

    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    pd.add_page_field(paragraph)


def set_cell_text(
    cell,
    text: str,
    *,
    size: float = 9.2,
    bold: bool = False,
    color: RGBColor | None = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    pd.set_run_font(
        run,
        size=size,
        color=color or pd.INK,
        bold=bold,
    )


def add_table(
    doc: Document,
    headers: tuple[str, ...],
    rows: Iterable[tuple[str, ...]],
    weights: tuple[float, ...],
    *,
    font_size: float = 8.8,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    header = table.rows[0]
    pd.mark_repeat_header(header)
    pd.prevent_row_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        pd.set_cell_shading(cell, "FCE9F1")
        pd.set_cell_border(cell, color="E5B7C8")
        set_cell_text(
            cell,
            value,
            size=font_size,
            bold=True,
            color=pd.BRAND_DARK,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row_values in rows:
        row = table.add_row()
        pd.prevent_row_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            pd.set_cell_shading(cell, pd.WHITE)
            pd.set_cell_border(cell, color="DEDDE2")
            set_cell_text(
                cell,
                value,
                size=font_size,
                bold=index == 0,
                color=pd.BRAND_DARK if index == 0 else pd.INK,
                align=(
                    WD_ALIGN_PARAGRAPH.CENTER
                    if index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
            )

    pd.apply_table_geometry(
        table,
        pd.column_widths_from_weights(weights, pd.TABLE_WIDTH_DXA),
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    return table


def add_title_block(doc: Document, fingerprint: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("APP 1.0 · 客户需求冻结资料")
    pd.set_run_font(run, size=10.5, color=pd.BRAND, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("MeiGallery App 1.0")
    pd.set_run_font(run, size=29, color=pd.BRAND_DARK, bold=True)
    run.add_break()
    run = paragraph.add_run("需求冻结确认单")
    pd.set_run_font(run, size=29, color=pd.BRAND_DARK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(
        "用于确认产品范围、服务承诺、页面交互、视觉成熟度和冻结生效条件"
    )
    pd.set_run_font(run, size=13, color=pd.MUTED)

    add_table(
        doc,
        ("项目", "内容", "项目", "内容"),
        (
            ("产品", "MeiGallery App", "App 版本", "1.0"),
            ("状态", "历史快照", "日期", "2026-07-30"),
            ("客户决策", "8 项", "专业门禁", "7 组"),
        ),
        (0.85, 2.25, 0.95, 2.45),
        font_size=9.1,
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("当前结论")
    pd.set_run_font(run, size=11, color=pd.BRAND, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    pd.add_inline_content(
        paragraph,
        "产品范围、功能和页面交互已具备客户确认条件；"
        "Figma 99 页/402 状态最终设计已完成；当前仍有 8 项客户决策、"
        "7 组专业门禁和客户视觉签署，"
        "因此尚未授权开发。",
        base_size=10.5,
        base_color=pd.BRAND_DARK,
        base_bold=True,
    )
    paragraph._p.get_or_add_pPr().append(
        OxmlElement("w:shd")
    )
    paragraph._p.pPr[-1].set(qn("w:fill"), "FFF3F7")

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    pd.add_inline_content(
        paragraph,
        "当前实时设计事实为 99 页/408 个正式状态，移动端 50 页/208 状态、后台 49 页/200 状态；"
        "本确认单中的 402/202 与 3,571 个动作是 APP-SET-08 六态增量前历史基线。",
        base_size=9.2,
        base_color=pd.MUTED,
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("基线指纹：")
    pd.set_run_font(run, size=8.6, color=pd.MUTED, bold=True)
    run = paragraph.add_run(fingerprint)
    pd.set_run_font(
        run,
        latin=pd.MONO_FONT,
        east_asia=pd.EAST_ASIA_FONT,
        size=7.6,
        color=pd.MUTED,
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(20)
    paragraph.paragraph_format.space_after = Pt(0)
    pd.add_inline_content(
        paragraph,
        "填写方式｜先完成第 3 章 8 项客户决策，再核对第 4 章专业门禁，"
        "最后在第 8 章选择确认结论并签署。",
        base_size=10,
        base_color=pd.INK,
        base_bold=True,
    )


def add_status_summary(doc: Document) -> None:
    heading = doc.add_paragraph("1. 冻结准备结论", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "本确认单是两份完整客户文档的签署入口，不替代产品需求确认书和"
        "逐页交互设计确认册，也不替代法务、安全、财务、运营或技术负责人的专业结论。",
        base_size=10.3,
    )
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "本次基线覆盖 99 个 Page ID、402 个 Figma 最终状态、99 个流程预览和 "
        "3,571 个有效交互动作；客户文档映射基线 179/179 个确定性图片映射已经建立。",
        base_size=10,
        base_color=pd.BRAND_DARK,
        base_bold=True,
    )
    add_table(
        doc,
        ("对象", "当前证据", "当前结论", "冻结条件"),
        (
            (
                "产品范围",
                "客户 PRD 与 C-01 至 C-08",
                "可确认，未冻结",
                "8 项客户决策形成书面结论",
            ),
            (
                "功能需求",
                "41 个需求编号、99 个追踪键",
                "可确认，未冻结",
                "调整同步到开发规格与追踪矩阵",
            ),
            (
                "页面交互",
                "50 个移动端页、49 个后台页",
                "可逐页确认",
                "逐页意见关闭或形成明确例外",
            ),
            (
                "原型映射",
                "156 张基础 + 23 张逐状态导出",
                "179/179 已建立",
                "映射与需求继续保持一致",
            ),
            (
                "像素级视觉",
                "99 页 / 402 状态最终稿",
                "已完成，待签署",
                "按 Page ID 关闭客户与设计意见",
            ),
            (
                "开发授权",
                "尚未创建 App 工程/API/migration",
                "未生效",
                "客户签署并关闭阻塞门禁",
            ),
        ),
        (1.05, 2.05, 1.25, 2.15),
        font_size=8.6,
    )


def add_decisions(doc: Document) -> None:
    heading = doc.add_paragraph("2. 客户需确认的 8 项决策", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "每项必须选择“接受”“调整”或“暂缓”。选择调整时填写替代结论；"
        "选择暂缓时，相关工程在最晚关闭点后不得继续。",
        base_size=10,
        base_color=pd.BRAND_DARK,
        base_bold=True,
    )

    for index, item in enumerate(CUSTOMER_DECISIONS):
        if index and index % 2 == 0:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.page_break_before = True
            spacer.paragraph_format.space_after = Pt(0)
        title = doc.add_paragraph(
            f"{item.decision_id}　{item.title}",
            style="Heading 2",
        )
        if index % 2 == 0:
            title.paragraph_format.space_before = Pt(0)
        add_table(
            doc,
            ("字段", "内容"),
            (
                ("建议基线", item.recommendation),
                ("影响", item.impact),
                ("最晚关闭点", item.deadline),
                ("来源", item.source_ids),
                ("客户选择", "□ 接受　　□ 调整　　□ 暂缓"),
                (
                    "调整意见",
                    "________________________________________________________",
                ),
            ),
            (1.18, 5.32),
            font_size=9.0,
        )


def add_visual_evidence(doc: Document) -> None:
    heading = doc.add_paragraph("3. 交互与视觉成熟度", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "99 页已经具备独立的信息结构、交互状态、业务规则和验收条件。"
        "Figma 最终文件完整覆盖移动端 50 页/202 状态、管理后台 49 页/200 状态，"
        "并完成 99 个流程预览；其中 3,571 个动作是 APP-SET-08 六态增量前历史基线，当前实时总数为 3,585；当前等待客户和设计负责人签署。",
        base_size=10.2,
    )

    pair = doc.add_table(rows=2, cols=2)
    pair.autofit = False
    pd.mark_repeat_header(pair.rows[0])
    pd.prevent_row_split(pair.rows[0])
    pd.prevent_row_split(pair.rows[1])
    for column_index, (alt, image_path) in enumerate(VISUAL_EVIDENCE[:2]):
        if not image_path.exists():
            raise FileNotFoundError(f"视觉证据不存在：{image_path}")
        caption_cell = pair.cell(0, column_index)
        pd.set_cell_shading(caption_cell, "FCE9F1")
        pd.set_cell_border(caption_cell, color="E8DCE1", size="2")
        set_cell_text(
            caption_cell,
            alt,
            size=8.8,
            bold=True,
            color=pd.BRAND_DARK,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        cell = pair.cell(1, column_index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        pd.set_cell_border(cell, color="E8DCE1", size="2")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(
            str(image_path),
            width=Inches(2.55),
        )
        pd.set_image_alt(shape, alt)
    pd.apply_table_geometry(
        pair,
        pd.column_widths_from_weights((3.25, 3.25), pd.TABLE_WIDTH_DXA),
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run("后台高风险页面示例")
    pd.set_run_font(run, size=10, color=pd.BRAND, bold=True)
    alt, image_path = VISUAL_EVIDENCE[2]
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    shape = paragraph.add_run().add_picture(
        str(image_path),
        width=Inches(6.05),
    )
    pd.set_image_alt(shape, alt)
    caption = doc.add_paragraph(
        f"图：{alt}。管理后台视觉可以更克制，但认证范围、版本、审批和审计状态必须完整。",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = False

    heading = doc.add_paragraph("最终视觉确认顺序", style="Heading 2")
    steps = (
        "通过 40｜Delivery Index 按 Page ID 核对 99 页和 402 个状态。",
        "通过 30｜Prototype Flows 核对 99 个流程的主操作、异常状态和安全出口。",
        "意见必须引用 Page ID、状态和目标调整，不以截图位置替代设计定位。",
        "修改后重新执行文字、Icon、热区、溢出、交互目标和产品边界 QA。",
    )
    num_id = create_compact_numbering(doc, "decimal")
    for step in steps:
        paragraph = doc.add_paragraph()
        apply_compact_numbering(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(4)
        pd.add_inline_content(paragraph, step, base_size=9.6)


def create_compact_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = pd.next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    level.append(lvl_text)
    justify = OxmlElement("w:lvlJc")
    justify.set(qn("w:val"), "left")
    level.append(justify)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), pd.BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), pd.BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), pd.EAST_ASIA_FONT)
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num_id = pd.next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_compact_numbering(paragraph, num_id: int) -> None:
    pd.apply_numbering(paragraph, num_id)


def add_professional_gates(doc: Document) -> None:
    heading = doc.add_paragraph("4. 上线前专业门禁", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "客户签署产品范围不等同于替专业负责人作结论。"
        "下列门禁必须分别提供责任人、结论日期和证据链接；"
        "超过最晚关闭点仍未完成时，对应工程或发布必须停止。",
        base_size=10,
    )

    for index, gate in enumerate(PROFESSIONAL_GATES):
        if index and index % 2 == 0:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.page_break_before = True
            spacer.paragraph_format.space_after = Pt(0)
        title = doc.add_paragraph(
            f"{gate.gate_id}　{gate.title}",
            style="Heading 2",
        )
        if index % 2 == 0:
            title.paragraph_format.space_before = Pt(0)
        add_table(
            doc,
            ("字段", "内容"),
            (
                ("必须形成的结论", gate.conclusion),
                ("责任方", gate.owner),
                ("最晚关闭点", gate.deadline),
                ("来源", gate.source_ids),
                ("关闭状态", "□ 已关闭　□ 有条件关闭　□ 未关闭"),
                (
                    "证据/意见",
                    "________________________________________________________",
                ),
            ),
            (1.18, 5.32),
            font_size=9.0,
        )


def add_freeze_rules(doc: Document) -> None:
    heading = doc.add_paragraph("5. 冻结生效与变更控制", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "只有以下条件全部满足，文档状态才能改为“App 1.0 需求已冻结”。",
        base_size=10.2,
        base_color=pd.BRAND_DARK,
        base_bold=True,
    )
    conditions = (
        "C-01 至 C-07 已签署；C-08 已确定，或明确接受工作名继续使用及品牌工程阻塞。",
        "产品需求确认书总确认结论已签署。",
        "逐页确认册意见已关闭，或形成带 Page ID、状态、责任人与关闭日期的例外清单。",
        "所有调整已同步到 PRD、开发规格、决策基线、Feature PRD、追踪矩阵、页面目录、原型和 DOCX。",
        "99 页、402 个 Figma 最终状态、3,571 个有效动作和 179 个客户图片映射继续通过自动校验。",
        "阻塞下一阶段的 G-01 至 G-07 已按最晚关闭点完成。",
        "产品 Owner、客户代表、设计负责人和技术负责人共同确认冻结结论。",
    )
    check_id = create_compact_numbering(doc, "bullet")
    for condition in conditions:
        paragraph = doc.add_paragraph()
        apply_compact_numbering(paragraph, check_id)
        pd.add_inline_content(
            paragraph,
            condition,
            base_size=9.7,
        )

    subheading = doc.add_paragraph("冻结后的变更要求", style="Heading 2")
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "App 与文档版本继续保持 1.0，不因讨论次数增加文档小版本。"
        "任何实质变化必须先记录原因、受影响需求/Page ID/状态、API 与数据影响、"
        "安全运营影响、工期、重新确认范围、生效时间和回滚条件；"
        "不得由设计、开发或测试人员自行选择。",
        base_size=10,
    )

    add_table(
        doc,
        ("变更类型", "处理方式", "是否重新确认"),
        (
            (
                "文字/排版修正",
                "直接更新，保持语义与追踪键不变",
                "否",
            ),
            (
                "业务规则、权限或状态变化",
                "完成影响评估并同步全部上下游",
                "是",
            ),
            (
                "新增页面或原生能力",
                "评估新 Page ID、API、最低 App 版本和发版",
                "是",
            ),
            (
                "未来支付、推送、礼物或装扮",
                "独立立项并正常升级 App",
                "是",
            ),
        ),
        (1.35, 3.85, 1.3),
        font_size=9.0,
    )


def add_fingerprint_section(
    doc: Document,
    fingerprints: list[tuple[str, str]],
    fingerprint: str,
) -> None:
    heading = doc.add_paragraph("6. 待确认冻结包", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "以下文件共同构成本次待确认基线。任一文件内容变化后都必须重新生成确认单；"
        "旧签署件不得自动覆盖新内容。完整 SHA-256 记录在同名 Markdown 清单中。",
        base_size=10,
    )
    rows = []
    for path, digest in fingerprints:
        label = Path(path).name
        rows.append((label, digest[:12], "待确认基线"))
    add_table(
        doc,
        ("文件", "短指纹", "状态"),
        tuple(rows),
        (4.25, 1.2, 1.05),
        font_size=8.7,
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("组合基线指纹：")
    pd.set_run_font(run, size=8.8, color=pd.MUTED, bold=True)
    run = paragraph.add_run(fingerprint)
    pd.set_run_font(
        run,
        latin=pd.MONO_FONT,
        east_asia=pd.EAST_ASIA_FONT,
        size=7.5,
        color=pd.MUTED,
    )


def add_signoff(doc: Document) -> None:
    heading = doc.add_paragraph("7. 客户确认与签署", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    paragraph = doc.add_paragraph()
    pd.add_inline_content(
        paragraph,
        "请选择一项总确认结论。若选择“原则同意”或“视觉需修改”，"
        "必须附上具体 C/G/Page ID 和调整意见。",
        base_size=10,
        base_color=pd.BRAND_DARK,
        base_bold=True,
    )
    options = (
        "□ 同意 C-01 至 C-08 的建议基线，进入需求冻结。",
        "□ 原则同意，但需按本确认单的调整意见同步后再冻结。",
        "□ 只同意产品范围与功能交互；Figma 最终视觉需按所列 Page ID 修改后再次确认。",
        "□ 暂不确认，需要重新讨论产品定位、范围或服务承诺。",
    )
    for option in options:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(option)
        pd.set_run_font(run, size=11, color=pd.INK, bold=True)

    add_table(
        doc,
        ("签署字段", "填写"),
        (
            ("客户/项目方", "____________________________________________"),
            ("确认人", "____________________________________________"),
            ("签字/盖章", "____________________________________________"),
            ("确认日期", "________年____月____日"),
            ("总体意见", "____________________________________________\n\n"),
        ),
        (1.55, 4.95),
        font_size=10,
    )

    subheading = doc.add_paragraph("签署后的执行顺序", style="Heading 2")
    steps = (
        "同步客户选择和调整到全部上游与下游文档。",
        "重新生成两份完整客户 DOCX、原型清单和本确认单。",
        "完成需求、Page ID、原型、DOCX、无障碍和全页渲染校验。",
        "由 Owner 记录冻结状态、组合指纹、签署人和日期。",
        "关闭视觉意见与阻塞门禁后，再冻结 API/DTO 契约并启动 KMP。",
    )
    num_id = create_compact_numbering(doc, "decimal")
    for step in steps:
        paragraph = doc.add_paragraph()
        apply_compact_numbering(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(2)
        pd.add_inline_content(paragraph, step, base_size=9.5)


def build_docx(
    fingerprints: list[tuple[str, str]],
    fingerprint: str,
) -> None:
    doc = Document()
    configure_compact_styles(doc)
    section = doc.sections[0]
    pd.configure_page(section)
    configure_customer_header_footer(section)

    add_title_block(doc, fingerprint)
    add_status_summary(doc)
    add_decisions(doc)
    add_visual_evidence(doc)
    add_professional_gates(doc)
    add_freeze_rules(doc)
    add_fingerprint_section(doc, fingerprints, fingerprint)
    add_signoff(doc)

    doc.core_properties.title = "MeiGallery App 1.0 需求冻结确认单"
    doc.core_properties.subject = (
        "客户产品范围、页面交互、视觉成熟度、专业门禁与冻结生效条件确认"
    )
    doc.core_properties.author = "MeiGallery 产品团队"
    doc.core_properties.keywords = (
        "MeiGallery, App 1.0, 需求冻结, 客户确认, 产品设计"
    )
    doc.core_properties.comments = (
        "版式：compact_reference_guide；首页：customer_pack；"
        "命名覆盖：MeiGallery 品牌粉色、Hiragino Sans GB；"
        "客户决策 8 项、专业门禁 7 组；"
        "当前状态为冻结准备中，不代表已授权开发。"
    )
    doc.core_properties.last_modified_by = "MeiGallery 产品团队"

    DOCX_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUTPUT)


def main() -> None:
    load_manifest()
    fingerprints = source_fingerprints()
    fingerprint = combined_fingerprint(fingerprints)
    generate_markdown(fingerprints, fingerprint)
    build_docx(fingerprints, fingerprint)
    print(f"已生成：{MD_OUTPUT}")
    print(f"已生成：{DOCX_OUTPUT}")
    print(f"组合基线指纹：{fingerprint}")


if __name__ == "__main__":
    main()
